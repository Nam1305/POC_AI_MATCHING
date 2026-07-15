"""
Stage 1 — Document Processing (Trích xuất văn bản từ file CV)

File này chịu trách nhiệm trích xuất văn bản sạch (clean text) từ file
PDF/DOCX nhận vào dưới dạng bytes (từ FastAPI UploadFile), để văn bản này
được đưa tiếp sang Stage 2 (parser.py) trích xuất thông tin có cấu trúc.

Pipeline xử lý:
  1. Trích xuất bằng PyMuPDF với logic nhận diện layout thông minh
     (phát hiện 1 cột hay 2 cột)
  2. Chấm điểm chất lượng văn bản trích xuất được (thang 0–100, theo heuristic)
  3. Nếu là PDF và điểm chất lượng < 60 → fallback sang OCR (Tesseract),
     thường xảy ra với PDF dạng scan/ảnh (không có text layer)
  4. Nếu là DOCX → trích xuất bằng python-docx (đọc từng đoạn văn/paragraph)
"""

from __future__ import annotations

import io
import re

import fitz                         # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document           # python-docx


# ---------------------------------------------------------------------------
# Public entry point — hàm duy nhất mà các module khác cần gọi
# ---------------------------------------------------------------------------

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Điểm vào chính của module: nhận bytes của file + tên file, tự động phân
    loại theo phần mở rộng (.pdf / .docx) rồi gọi hàm trích xuất tương ứng.

    Với PDF: trích xuất bằng smart layout trước, nếu chất lượng quá thấp
    (score < 60, khả năng cao là PDF dạng scan) thì tự động fallback sang OCR.

    Trả về text đã được làm sạch (chuẩn hóa khoảng trắng, loại bỏ ký tự
    điều khiển — control characters).

    Raise ValueError nếu phần mở rộng file không được hỗ trợ (chỉ nhận
    .pdf / .docx).
    """
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        text = extract_text_smart_layout(file_bytes)
        quality = evaluate_extracted_text_quality(text)
        if quality["score"] < 60:
            # Fallback: rasterize + OCR
            text = ocr_pdf_with_tesseract(file_bytes)
    elif ext in ("docx", "doc"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext} (only .pdf / .docx allowed)")

    return clean_text(text)


# ---------------------------------------------------------------------------
# PDF — smart layout extraction (port from POC)
# ---------------------------------------------------------------------------

def extract_text_smart_layout(file_bytes: bytes) -> str:
    """
    Trích xuất text từ PDF ở cấp độ block bằng PyMuPDF, có nhận diện layout
    1 cột hay 2 cột cho từng trang.

    Vì sao cần: nhiều mẫu CV có layout 2 cột (ví dụ cột trái là timeline/kỹ
    năng, cột phải là nội dung chi tiết) — nếu đọc theo thứ tự tọa độ y
    thông thường (top-to-bottom) sẽ bị trộn lẫn nội dung 2 cột, sai thứ tự
    đọc. Hàm này xử lý bằng cách: đọc phần header trước, rồi đọc cột phải,
    rồi mới đến cột trái — để giữ đúng thứ tự đọc tự nhiên của CV.

    Cách phát hiện 2 cột: đếm số block có tâm (center_x) nằm bên trái 45%
    chiều rộng trang và số block nằm bên phải 55% chiều rộng trang; nếu cả
    2 phía đều có ít nhất 2 block thì coi là layout 2 cột.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    all_pages: list[str] = []

    for page in doc:
        page_width  = page.rect.width
        page_height = page.rect.height
        raw_blocks  = page.get_text("blocks")

        blocks = []
        for block in raw_blocks:
            x0, y0, x1, y1, text, *_ = block
            text = text.strip()
            if not text or len(text) < 2:
                continue
            blocks.append({
                "x0": x0, "y0": y0, "x1": x1, "y1": y1,
                "center_x": (x0 + x1) / 2,
                "text": text,
            })

        if not blocks:
            continue

        blocks.sort(key=lambda b: (b["y0"], b["x0"]))
        left_count  = sum(1 for b in blocks if b["center_x"] < page_width * 0.45)
        right_count = sum(1 for b in blocks if b["center_x"] > page_width * 0.55)
        is_two_col  = left_count >= 2 and right_count >= 2

        parts: list[str] = []
        if not is_two_col:
            parts.extend(b["text"] for b in blocks)
        else:
            header_limit = page_height * 0.16
            header_b = sorted(
                [b for b in blocks if b["y0"] < header_limit],
                key=lambda b: (b["y0"], b["x0"]),
            )
            body_b  = [b for b in blocks if b["y0"] >= header_limit]
            right_b = sorted(
                [b for b in body_b if b["center_x"] >= page_width * 0.45],
                key=lambda b: (b["y0"], b["x0"]),
            )
            left_b  = sorted(
                [b for b in body_b if b["center_x"] < page_width * 0.45],
                key=lambda b: (b["y0"], b["x0"]),
            )
            parts.extend(b["text"] for b in header_b)
            parts.extend(b["text"] for b in right_b)
            parts.extend(b["text"] for b in left_b)

        all_pages.append("\n".join(parts))

    doc.close()
    return "\n\n".join(all_pages)


# ---------------------------------------------------------------------------
# Quality heuristic — ước lượng độ tin cậy của text vừa trích xuất
# ---------------------------------------------------------------------------

def evaluate_extracted_text_quality(text: str) -> dict:
    """
    Chấm điểm heuristic 0–100 cho chất lượng text vừa trích xuất được.
    Điểm càng thấp thì khả năng đây là PDF dạng scan/ảnh (không có text
    layer thật) càng cao → cần OCR fallback.

    Các tín hiệu (signal) được dùng để trừ điểm:
      - Độ dài text quá ngắn (< 100 ký tự) → trích xuất kém
      - Số lượng từ quá ít (< 30 từ)       → nội dung không đáng kể
      - Tỷ lệ ký tự rác (garbage char, ví dụ ký tự thay thế "�") quá cao
      - Độ dài từ trung bình bất thường (từ có nghĩa thường dài 4–8 ký tự;
        quá ngắn hoặc quá dài gợi ý lỗi encode/OCR)

    Trả về dict {"score": điểm cuối (đã clamp >= 0), "reasons": danh sách lý
    do bị trừ điểm, "word_count": số từ đếm được}.
    """
    score = 100
    reasons: list[str] = []

    if len(text) < 100:
        score -= 60
        reasons.append("text too short")

    words = re.findall(r"\b\w+\b", text)
    if len(words) < 30:
        score -= 30
        reasons.append("too few words")

    if text:
        garbage = sum(1 for c in text if c in "�")
        if garbage / max(len(text), 1) > 0.02:
            score -= 20
            reasons.append("high garbage char ratio")

    if words:
        avg_len = sum(len(w) for w in words) / len(words)
        if avg_len < 2 or avg_len > 15:
            score -= 15
            reasons.append(f"abnormal avg word length {avg_len:.1f}")

    return {"score": max(0, score), "reasons": reasons, "word_count": len(words)}


# ---------------------------------------------------------------------------
# OCR fallback — dùng khi PDF là bản scan/ảnh, không có text layer tốt
# ---------------------------------------------------------------------------

def ocr_pdf_with_tesseract(file_bytes: bytes, dpi: int = 200, lang: str = "eng+vie") -> str:
    """
    Render (rasterize) từng trang PDF thành ảnh PNG ở độ phân giải `dpi`,
    sau đó chạy Tesseract OCR để đọc chữ từ ảnh (dùng khi PDF không có text
    layer đáng tin cậy — ví dụ file scan).

    Yêu cầu: máy chạy phải cài binary Tesseract (trong Docker: cần cài gói
    tesseract-ocr + tesseract-ocr-vie để đọc được cả tiếng Anh và tiếng Việt).
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text: list[str] = []

    for page in doc:
        # Render page to PNG bytes
        zoom = dpi / 72                  # 72 DPI is PyMuPDF default
        matrix = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))

        page_text = pytesseract.image_to_string(img, lang=lang)
        pages_text.append(page_text)

    doc.close()
    return "\n\n".join(pages_text)


# ---------------------------------------------------------------------------
# DOCX — trích xuất text từ file Word
# ---------------------------------------------------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Đọc toàn bộ đoạn văn (paragraphs) và nội dung ô bảng (table cells) từ
    file .docx, nối lại thành 1 chuỗi text (mỗi phần tử 1 dòng).
    """
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Text cleaning — chuẩn hóa text trước khi trả về cho tầng sau
# ---------------------------------------------------------------------------

def clean_text(text: str) -> str:
    """
    Chuẩn hóa khoảng trắng và loại bỏ ký tự điều khiển (control characters):
      - Thay ký tự null (\\x00) bằng khoảng trắng
      - Gộp nhiều space/tab liên tiếp thành 1 space
      - Gộp 3+ dòng trống liên tiếp thành tối đa 2 dòng trống (1 dòng trắng phân cách)
      - Cắt khoảng trắng thừa ở đầu/cuối chuỗi
    """
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
